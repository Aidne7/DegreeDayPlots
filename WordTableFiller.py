from docx import Document
import pyautogui
import os
import time
import pygetwindow as gw
from PIL import Image
import pandas as pd


# Function to fill the word table using the file dd_word.docx
def fillWordTab(ddTable):
    columns_to_drop = [4, 8]  # 0-based index of columns to drop
    wordDF = ddTable.drop(ddTable.columns[columns_to_drop], axis=1)
    # Calculate the average of each column
    avg_row = wordDF.mean().to_frame().T

    # Add the average row to the DataFrame
    wordDF = pd.concat([wordDF, avg_row], ignore_index=True)

    # Rounds the numbers to the nearest integer, converts them to interger, converts them to strings
    wordDF = wordDF.round(0).astype(int).astype(str)

    # Adds the word "State" to the bottom cell of the first column
    wordDF.iloc[-1, 0] = "State"

    # Opens the Word doc
    doc = Document('dd_word.docx')

    # Access the existing table (assuming it's the first table in the document)
    table = doc.tables[0]

    # Fill the existing table with data from the DataFrame
    for i, row in wordDF.iterrows():
        for j, col in enumerate(wordDF.columns):
            cell = table.cell(int(i) + 1, j)
            cell.text = str(row[col])
            # Center align the text in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = False  # Example: Keep the text unbolded
                    paragraph.alignment = 1  # 0=left, 1=center, 2=right

    # Save the document
    doc.save('dd_word_filled.docx')

    doc_path = "dd_word_filled.docx"

    # Open the document
    os.startfile(doc_path)

    # Wait for the application to open
    time.sleep(3)

    # Find the Word window
    word_window = gw.getWindowsWithTitle("dd_word_filled")[0]

    # Maximize the Word window
    word_window.maximize()
    time.sleep(1)

    # Take a screenshot of the entire screen
    screenshot = pyautogui.screenshot()

    # Save the screenshot to a file
    screenshot.save('screenshot.png')
    time.sleep(1)
    screenshot = Image.open('screenshot.png')

    # Define the region to crop (left, top, right, bottom)
    region = (577, 342, 1323, 693)

    # Crop the image to the specified region
    cropped_image = screenshot.crop(region)

    # Save the cropped image
    cropped_image.save('DDTable____.png')

    # Close the Word application
    word_window.close()
    time.sleep(1)

    # Closes the Word file
    os.remove("screenshot.png")
    os.remove("dd_word_filled.docx")
